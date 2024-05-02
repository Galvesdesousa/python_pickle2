import pickle
import csv
import os
from docx import Document
import platform
from openpyxl import Workbook

def clear_screen():
    if platform.system() == 'Windows':
        os.system('cls')
    else:
        os.system('clear')

def salvar_contatos(contatos):
    if os.path.exists('contatos.pickle'):
        limpar_arquivo()
    with open('contatos.pickle', 'wb') as arquivo:
        pickle.dump(contatos, arquivo)
    print("Contatos salvos com sucesso!")

def carregar_contatos():
    try:
        with open('contatos.pickle', 'rb') as arquivo:
            contatos = pickle.load(arquivo)
        print("Contatos carregados com sucesso!")
        return contatos
    except FileNotFoundError:
        print("Arquivo de contatos não encontrado. Criando nova lista.")
        return []

def adicionar_contato(contatos):
    nome = input("Digite o nome do contato: ")
    email = input("Digite o email do contato: ")
    telefone = input("Digite o telefone do contato: ")
    contatos.append({"nome": nome, "email": email, "telefone": telefone})
    print("Contato adicionado com sucesso!")

def listar_contatos(contatos):
    if contatos:
        print("Lista de contatos:")
        for contato in contatos:
            print(f"Nome: {contato['nome']}, Email: {contato['email']}, Telefone: {contato['telefone']}")
    else:
        print("Nenhum contato na lista.")

def limpar_arquivo():
    with open('contatos.pickle', 'wb') as arquivo:
        arquivo.truncate(0)
    print("Arquivo limpo com sucesso!")

def eliminar_arquivo():
    if os.path.exists('contatos.pickle'):
        os.remove('contatos.pickle')
        print("Arquivo eliminado com sucesso!")
    else:
        print("Arquivo não encontrado.")

def exportar_para_csv(contatos):
    with open('contatos.csv', 'w', newline='') as arquivo_csv:
        escritor = csv.DictWriter(arquivo_csv, fieldnames=['nome', 'email', 'telefone'])
        escritor.writeheader()
        for contato in contatos:
            escritor.writerow(contato)
    print("Contatos exportados para CSV com sucesso!")

def exportar_para_word(contatos):
    doc = Document()
    doc.add_heading('Lista de Contatos', level=1)
    for contato in contatos:
        doc.add_paragraph(f"Nome: {contato['nome']}, Email: {contato['email']}, Telefone: {contato['telefone']}")
    doc.save('contatos.docx')
    print("Contatos exportados para Word com sucesso!")

def exportar_para_excel(contatos):
    wb = Workbook()
    ws = wb.active
    ws.append(['Nome', 'Email', 'Telefone'])
    for contato in contatos:
        ws.append([contato['nome'], contato['email'], contato['telefone']])
    wb.save('contatos.xlsx')
    print("Contatos exportados para Excel com sucesso!")

def menu():
    print("\n### Menu ###")
    print("1. Adicionar Contato")
    print("2. Listar Contatos")
    print("3. Salvar Contatos")
    print("4. Carregar Contatos")
    print("5. Limpar Arquivo")
    print("6. Eliminar Arquivo")
    print("7. Exportar para CSV")
    print("8. Exportar para Word")
    print("9. Exportar para Excel")
    print("10. Sair")

    escolha = input("Escolha uma opção: ")
    return escolha

if __name__ == "__main__":
    contatos = []
    while True:
        clear_screen()
        escolha = menu()
        if escolha == "1":
            adicionar_contato(contatos)
        elif escolha == "2":
            listar_contatos(contatos)
        elif escolha == "3":
            salvar_contatos(contatos)
        elif escolha == "4":
            contatos = carregar_contatos()
        elif escolha == "5":
            limpar_arquivo()
        elif escolha == "6":
            eliminar_arquivo()
        elif escolha == "7":
            exportar_para_csv(contatos)
        elif escolha == "8":
            exportar_para_word(contatos)
        elif escolha == "9":
            exportar_para_excel(contatos)
        elif escolha == "10":
            print("Saindo...")
            break
        else:
            print("Escolha inválida. Tente novamente.")
        input("Pressione Enter para continuar...")
